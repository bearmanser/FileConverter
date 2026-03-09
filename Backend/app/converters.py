from __future__ import annotations

import csv
import html
import json
import re
import textwrap
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from markdown import markdown
from openpyxl import Workbook, load_workbook
from PIL import Image
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from .conversion_map import (
    DOCUMENT_FORMATS,
    IMAGE_FORMATS,
    MEDIA_TYPES,
    SUPPORTED_CONVERSIONS,
    TABULAR_FORMATS,
)


class ConversionError(Exception):
    """Raised when the requested conversion cannot be completed."""


def normalize_extension(value: str) -> str:
    return value.lower().strip().lstrip(".")


def get_output_media_type(target_format: str) -> str:
    return MEDIA_TYPES.get(target_format, "application/octet-stream")


def build_download_name(filename: str, target_format: str) -> str:
    stem = Path(filename).stem or "converted-file"
    safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "-", stem).strip("-") or "converted-file"
    return f"{safe_stem}.{target_format}"


def convert_file(
    *,
    filename: str,
    content: bytes,
    source_format: str,
    target_format: str,
) -> tuple[bytes, str]:
    source_format = normalize_extension(source_format)
    target_format = normalize_extension(target_format)

    allowed_targets = SUPPORTED_CONVERSIONS.get(source_format, [])
    if target_format not in allowed_targets:
        raise ConversionError(
            f"Conversion from {source_format or 'unknown'} to {target_format or 'unknown'} is not supported."
        )

    if source_format in DOCUMENT_FORMATS:
        converted = convert_document(
            content=content,
            source_format=source_format,
            target_format=target_format,
            filename=filename,
        )
    elif source_format in IMAGE_FORMATS:
        converted = convert_image(
            content=content,
            source_format=source_format,
            target_format=target_format,
        )
    elif source_format in TABULAR_FORMATS:
        converted = convert_tabular(
            content=content,
            source_format=source_format,
            target_format=target_format,
            filename=filename,
        )
    else:
        raise ConversionError(f"Files with the .{source_format} extension are not supported yet.")

    return converted, get_output_media_type(target_format)


def convert_document(
    *,
    content: bytes,
    source_format: str,
    target_format: str,
    filename: str,
) -> bytes:
    title = Path(filename).stem or "Converted document"

    if target_format == "txt":
        return document_to_text(content=content, source_format=source_format).encode("utf-8")

    if target_format == "html":
        if source_format == "html":
            return content
        if source_format == "md":
            return markdown_to_html_document(decode_text(content), title).encode("utf-8")
        text = document_to_text(content=content, source_format=source_format)
        return plain_text_to_html_document(text, title).encode("utf-8")

    if target_format == "md":
        if source_format == "md":
            return content
        if source_format == "html":
            markdown_text = html_to_markdown(decode_text(content))
        else:
            markdown_text = document_to_text(content=content, source_format=source_format)
        return normalize_text(markdown_text).encode("utf-8")

    plain_text = document_to_text(content=content, source_format=source_format)

    if target_format == "docx":
        return plain_text_to_docx(plain_text, title)

    if target_format == "pdf":
        return plain_text_to_pdf(plain_text, title)

    raise ConversionError(f"Unsupported document target format: {target_format}")


def document_to_text(*, content: bytes, source_format: str) -> str:
    if source_format == "txt":
        return normalize_text(decode_text(content))
    if source_format == "md":
        markdown_html = markdown(decode_text(content), extensions=["extra", "tables", "nl2br"])
        return html_to_text(markdown_html)
    if source_format == "html":
        return html_to_text(decode_text(content))
    if source_format == "docx":
        return docx_to_text(content)
    if source_format == "pdf":
        return pdf_to_text(content)
    raise ConversionError(f"Unsupported document source format: {source_format}")


def decode_text(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def normalize_text(value: str) -> str:
    text = value.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def html_to_text(raw_html: str) -> str:
    soup = BeautifulSoup(raw_html, "html.parser")
    for element in soup(["script", "style"]):
        element.decompose()
    return normalize_text(soup.get_text("\n"))


def html_to_markdown(raw_html: str) -> str:
    soup = BeautifulSoup(raw_html, "html.parser")
    root = soup.body or soup
    blocks: list[str] = []

    for child in root.children:
        rendered = render_markdown_block(child)
        if rendered:
            blocks.append(rendered)

    if not blocks:
        return normalize_text(root.get_text("\n"))

    return normalize_text("\n\n".join(blocks))


def render_markdown_block(node: Any) -> str:
    if isinstance(node, NavigableString):
        return normalize_text(str(node))

    if not isinstance(node, Tag):
        return ""

    name = node.name.lower()
    text = normalize_text(node.get_text(" ", strip=True))

    if not text and name not in {"ul", "ol"}:
        return ""

    if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        return f"{'#' * int(name[1])} {text}"
    if name == "p":
        return text
    if name == "blockquote":
        return "\n".join(f"> {line}" for line in text.splitlines())
    if name == "pre":
        return f"```\n{node.get_text()}\n```"
    if name == "ul":
        items = [normalize_text(item.get_text(" ", strip=True)) for item in node.find_all("li", recursive=False)]
        return "\n".join(f"- {item}" for item in items if item)
    if name == "ol":
        items = [normalize_text(item.get_text(" ", strip=True)) for item in node.find_all("li", recursive=False)]
        return "\n".join(f"{index}. {item}" for index, item in enumerate(items, start=1) if item)

    nested_blocks = [render_markdown_block(child) for child in node.children]
    return normalize_text("\n\n".join(block for block in nested_blocks if block))


def markdown_to_html_document(markdown_text: str, title: str) -> str:
    body = markdown(markdown_text, extensions=["extra", "tables", "nl2br"])
    escaped_title = html.escape(title)
    return (
        "<!doctype html>\n"
        "<html lang=\"en\">\n"
        "<head>\n"
        "  <meta charset=\"utf-8\" />\n"
        f"  <title>{escaped_title}</title>\n"
        "  <style>\n"
        "    body { font-family: Arial, sans-serif; max-width: 760px; margin: 40px auto; line-height: 1.6; }\n"
        "    pre { background: #f4f4f4; padding: 12px; overflow-x: auto; }\n"
        "    table { border-collapse: collapse; width: 100%; }\n"
        "    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }\n"
        "  </style>\n"
        "</head>\n"
        f"<body>\n{body}\n</body>\n"
        "</html>\n"
    )


def plain_text_to_html_document(text: str, title: str) -> str:
    escaped_title = html.escape(title)
    paragraphs = [
        f"<p>{html.escape(paragraph).replace(chr(10), '<br />')}</p>"
        for paragraph in normalize_text(text).split("\n\n")
        if paragraph.strip()
    ]
    body = "\n".join(paragraphs) or "<p></p>"
    return (
        "<!doctype html>\n"
        "<html lang=\"en\">\n"
        "<head>\n"
        "  <meta charset=\"utf-8\" />\n"
        f"  <title>{escaped_title}</title>\n"
        "  <style>body { font-family: Arial, sans-serif; max-width: 760px; margin: 40px auto; line-height: 1.6; }</style>\n"
        "</head>\n"
        f"<body>\n{body}\n</body>\n"
        "</html>\n"
    )


def docx_to_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return normalize_text("\n\n".join(paragraphs))


def pdf_to_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(page for page in pages if page)
    return normalize_text(text)


def plain_text_to_docx(text: str, title: str) -> bytes:
    document = Document()
    document.core_properties.title = title

    normalized = normalize_text(text)
    if not normalized:
        document.add_paragraph("")
    else:
        for paragraph in normalized.split("\n\n"):
            document.add_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def plain_text_to_pdf(text: str, title: str) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.setTitle(title)

    page_width, page_height = letter
    margin = 48
    line_height = 14
    wrap_width = 95
    y = page_height - margin

    pdf.setFont("Helvetica", 11)

    normalized = normalize_text(text)
    paragraphs = normalized.split("\n\n") if normalized else [""]

    for paragraph in paragraphs:
        lines = textwrap.wrap(paragraph, width=wrap_width) or [""]
        for line in lines:
            if y <= margin:
                pdf.showPage()
                pdf.setFont("Helvetica", 11)
                y = page_height - margin
            pdf.drawString(margin, y, line)
            y -= line_height
        y -= line_height // 2

    pdf.save()
    return buffer.getvalue()


def convert_image(*, content: bytes, source_format: str, target_format: str) -> bytes:
    with Image.open(BytesIO(content)) as image:
        output = BytesIO()

        if target_format == "pdf":
            pdf_ready = image.convert("RGB")
            pdf_ready.save(output, format="PDF")
            return output.getvalue()

        save_format = {
            "jpg": "JPEG",
            "jpeg": "JPEG",
            "png": "PNG",
            "webp": "WEBP",
        }[target_format]

        processed = image
        if target_format in {"jpg", "jpeg"}:
            background = Image.new("RGB", image.size, (255, 255, 255))
            if image.mode in ("RGBA", "LA"):
                background.paste(image, mask=image.getchannel("A"))
                processed = background
            else:
                processed = image.convert("RGB")
        elif target_format == "png":
            processed = image.convert("RGBA") if image.mode not in ("RGB", "RGBA") else image
        else:
            processed = image.convert("RGB") if image.mode not in ("RGB", "RGBA") else image

        processed.save(output, format=save_format)
        return output.getvalue()


def convert_tabular(
    *,
    content: bytes,
    source_format: str,
    target_format: str,
    filename: str,
) -> bytes:
    rows = parse_tabular(content=content, source_format=source_format)
    title = Path(filename).stem or "Converted table"

    if target_format == "csv":
        return rows_to_csv(rows).encode("utf-8")
    if target_format == "json":
        return json.dumps(rows, indent=2, ensure_ascii=True).encode("utf-8")
    if target_format == "txt":
        return rows_to_text(rows).encode("utf-8")
    if target_format == "xlsx":
        return rows_to_xlsx(rows)
    if target_format == "pdf":
        return plain_text_to_pdf(rows_to_text(rows), title)

    raise ConversionError(f"Unsupported tabular target format: {target_format}")


def parse_tabular(*, content: bytes, source_format: str) -> list[dict[str, Any]]:
    if source_format == "csv":
        return parse_csv(content)
    if source_format == "json":
        return parse_json(content)
    if source_format == "xlsx":
        return parse_xlsx(content)
    raise ConversionError(f"Unsupported tabular source format: {source_format}")


def parse_csv(content: bytes) -> list[dict[str, Any]]:
    text = decode_text(content)
    stream = StringIO(text)
    reader = csv.reader(stream)
    rows = list(reader)

    if not rows:
        return []

    header_row = rows[0]
    headers = [
        normalize_header(header, index)
        for index, header in enumerate(header_row, start=1)
    ]

    data_rows = rows[1:] or [[]]
    return [dict(zip(headers, pad_row(row, len(headers)), strict=False)) for row in data_rows]


def parse_json(content: bytes) -> list[dict[str, Any]]:
    payload = json.loads(decode_text(content))

    if isinstance(payload, list):
        if not payload:
            return []
        if all(isinstance(item, dict) for item in payload):
            return normalize_rows(payload)
        return [{"value": item} for item in payload]

    if isinstance(payload, dict):
        if all(not isinstance(value, (list, dict)) for value in payload.values()):
            return [payload]
        if isinstance(payload.get("rows"), list):
            return parse_json(json.dumps(payload["rows"]).encode("utf-8"))
        return [{"key": key, "value": json.dumps(value, ensure_ascii=True)} for key, value in payload.items()]

    return [{"value": payload}]


def parse_xlsx(content: bytes) -> list[dict[str, Any]]:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    worksheet = workbook.active
    row_values = list(worksheet.iter_rows(values_only=True))

    if not row_values:
        return []

    headers = [
        normalize_header(value, index)
        for index, value in enumerate(row_values[0], start=1)
    ]
    data_rows = row_values[1:] or [tuple("" for _ in headers)]

    return [
        dict(zip(headers, pad_row(list(row), len(headers)), strict=False))
        for row in data_rows
    ]


def normalize_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    headers = collect_headers(rows)
    return [
        {header: stringify_cell(row.get(header, "")) for header in headers}
        for row in rows
    ]


def collect_headers(rows: list[dict[str, Any]]) -> list[str]:
    headers: list[str] = []
    for row in rows:
        for key in row:
            if key not in headers:
                headers.append(str(key))
    return headers or ["value"]


def normalize_header(value: Any, index: int) -> str:
    header = str(value).strip() if value is not None else ""
    return header or f"column_{index}"


def pad_row(row: list[Any], size: int) -> list[str]:
    cells = [stringify_cell(cell) for cell in row]
    if len(cells) < size:
        cells.extend([""] * (size - len(cells)))
    return cells[:size]


def stringify_cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=True)
    return str(value)


def rows_to_csv(rows: list[dict[str, Any]]) -> str:
    headers = collect_headers(rows)
    stream = StringIO()
    writer = csv.DictWriter(stream, fieldnames=headers)
    writer.writeheader()

    for row in rows:
        writer.writerow({header: stringify_cell(row.get(header, "")) for header in headers})

    return stream.getvalue()


def rows_to_text(rows: list[dict[str, Any]]) -> str:
    headers = collect_headers(rows)
    widths = {
        header: max([len(header), *(len(stringify_cell(row.get(header, ""))) for row in rows)])
        for header in headers
    }

    header_line = " | ".join(header.ljust(widths[header]) for header in headers)
    divider_line = "-+-".join("-" * widths[header] for header in headers)
    body_lines = [
        " | ".join(stringify_cell(row.get(header, "")).ljust(widths[header]) for header in headers)
        for row in rows
    ]

    table_lines = [header_line, divider_line, *body_lines]
    return "\n".join(table_lines).rstrip()


def rows_to_xlsx(rows: list[dict[str, Any]]) -> bytes:
    headers = collect_headers(rows)
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Converted"
    worksheet.append(headers)

    for row in rows:
        worksheet.append([stringify_cell(row.get(header, "")) for header in headers])

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()
